"""
Unit tests for DOCXExtractor.
"""

from docx import Document

from athena.indexing.extractors.docx import DOCXExtractor


def test_supported_extensions():
    extractor = DOCXExtractor()

    assert extractor.supported_extensions == (".docx",)


def test_extract_docx(tmp_path):
    document_path = tmp_path / "sample.docx"

    document = Document()
    document.add_heading("Athena", level=1)
    document.add_paragraph("Offline AI Research Platform")
    document.save(document_path)

    extractor = DOCXExtractor()

    extracted = extractor.extract(document_path)

    assert extracted.document_id == "sample.docx"
    assert extracted.title == "sample"
    assert "Athena" in extracted.text
    assert "Offline AI Research Platform" in extracted.text
    assert extracted.page_count == 1
