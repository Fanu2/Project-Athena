"""
Microsoft Word (.docx) document extractor.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from athena.indexing.extractors.base import BaseExtractor
from athena.indexing.models import (
    ExtractedDocument,
    ExtractedPage,
)


class DOCXExtractor(BaseExtractor):
    """Extract text from Microsoft Word documents."""

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        """Return supported file extensions."""
        return (".docx",)

    def extract(
        self,
        document: Path,
    ) -> ExtractedDocument:
        """
        Extract text from a DOCX document.

        Args:
            document:
                Path to the DOCX document.

        Returns:
            ExtractedDocument.
        """

        doc = Document(str(document))

        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())

        return ExtractedDocument(
            document_id=document.name,
            path=document,
            title=document.stem,
            text=text,
            pages=(
                ExtractedPage(
                    page_number=1,
                    text=text,
                ),
            ),
            page_count=1,
        )
